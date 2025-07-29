from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 创建 Word 文档
doc = Document()

# 设置页面边距
for section in doc.sections:
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# 标题
heading = doc.add_heading('[Your Full Name]', 0)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('[Your City, China] | [Your Phone Number] | [Your Email Address] | [LinkedIn or GitHub URL]')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('[Optional: Personal Website or Portfolio URL]', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

# Objective
doc.add_heading('Objective', level=1)
doc.add_paragraph(
    'Highly motivated and adaptable Software Developer with experience in backend API development for web and mobile applications. '
    'Proficient in PHP, JavaScript, Golang, and Python, with a strong focus on building scalable and efficient web backend APIs. '
    'Seeking a remote Software Developer position to contribute technical expertise and deliver high-quality solutions in a dynamic, global team.'
)

# Skills
doc.add_heading('Skills', level=1)
skills = [
    'Programming Languages: Golang, Python, PHP, JavaScript',
    'Frameworks & Tools: Flask (Python), FastAPI (Python), Laravel (PHP), Node.js (JavaScript), Git',
    'Technologies: RESTful APIs, MySQL, PostgreSQL, Docker, Linux',
    'Concepts: Backend API Development, Microservices, Web Application Architecture, Agile Development',
    'Soft Skills: Problem-solving, Team Collaboration, Adaptability, Strong Communication'
]
for skill in skills:
    doc.add_paragraph(skill, style='List Bullet')

# Professional Experience
doc.add_heading('Professional Experience', level=1)
doc.add_heading('Backend Developer (Freelance/Self-Employed)', level=2)
doc.add_paragraph('[Your City, China] | [Month, Year] – Present')
experiences = [
    'Designed and developed RESTful APIs for web applications using Golang and Python (Flask, FastAPI), improving system scalability and performance.',
    'Built and maintained backend services for websites, focusing on secure data processing and efficient request handling.',
    'Integrated databases (MySQL, PostgreSQL) to support dynamic content delivery and user management.',
    'Collaborated with front-end developers to ensure seamless API integration for web applications.',
    'Managed version control using Git for efficient code deployment.'
]
for exp in experiences:
    doc.add_paragraph(exp, style='List Bullet')

doc.add_heading('Web and Mobile App Backend Developer', level=2)
doc.add_paragraph('[Company Name or "Various Projects"] | [Your City, China] | [Month, Year] – [Month, Year]')
experiences = [
    'Developed backend APIs for web and mobile applications using PHP (Laravel) and JavaScript (Node.js), enabling robust functionality for user authentication, data retrieval, and payment processing.',
    'Optimized API performance by implementing caching mechanisms and database query optimizations.',
    'Designed database schemas and managed migrations for MySQL databases to support application scalability.',
    'Worked in Agile teams, participating in sprint planning and code reviews to ensure high-quality deliverables.',
    'Debugged and resolved critical backend issues, improving system reliability and user experience.'
]
for exp in experiences:
    doc.add_paragraph(exp, style='List Bullet')

# Projects
doc.add_heading('Projects', level=1)
doc.add_heading('Personal Web API Project', level=2)
doc.add_paragraph('[Month, Year] – Present')
projects = [
    'Built a RESTful API using FastAPI (Python) for a personal web application, supporting file uploads and user authentication.',
    'Implemented SSH key-based Git workflows for secure code management on GitHub.',
    'Technologies: Python, FastAPI, PostgreSQL.'
]
for proj in projects:
    doc.add_paragraph(proj, style='List Bullet')

doc.add_heading('E-Commerce Backend API', level=2)
doc.add_paragraph('[Month, Year] – [Month, Year]')
projects = [
    'Developed a backend API for an e-commerce platform using PHP (Laravel), handling product listings, user carts, and order processing.',
    'Integrated third-party payment APIs and optimized database queries for high-traffic scenarios.',
    'Technologies: PHP, Laravel, MySQL, JavaScript.'
]
for proj in projects:
    doc.add_paragraph(proj, style='List Bullet')

# Education
doc.add_heading('Education', level=1)
doc.add_heading('Bachelor’s Degree in [Your Major, e.g., Computer Science]', level=2)
doc.add_paragraph('[Your University Name], [Your City, China] | [Year] – [Year]')
doc.add_paragraph('Relevant Coursework: Data Structures, Algorithms, Web Development, Database Systems', style='List Bullet')

# Certifications
doc.add_heading('Certifications', level=1)
doc.add_paragraph('[Optional: Add relevant certifications, e.g., Online Course in Golang]', style='List Bullet')
doc.add_paragraph('[Example: Completed "Python for Web Development" online course, [Platform], [Year]]', style='List Bullet')

# Additional Information
doc.add_heading('Additional Information', level=1)
info = [
    'Languages: Mandarin Chinese (Native), English (Proficient)',
    'Availability: Flexible for remote work across global time zones.',
    'Interests: Open-source contributions, learning new frameworks, cloud computing.'
]
for item in info:
    doc.add_paragraph(item, style='List Bullet')

# 设置字体
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)

# 保存文档
doc.save('resume.docx')